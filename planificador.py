import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENTATION
from datetime import date
import locale

# Configurar la p谩gina de Streamlit
st.set_page_config(page_title="Generador de Planificaciones", page_icon="")

try:
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
except locale.Error:
    st.warning("锔 Advertencia: No se pudo establecer la configuraci贸n regional 'es_ES.UTF-8'. Se usar谩 la configuraci贸n por defecto.")

st.title(" Generador de Planificaciones Diarias para Docentes")

# Inicializar session_state para almacenar planificaciones
if "planificaciones" not in st.session_state:
    st.session_state.planificaciones = []

# Campos para el encabezado
semana = st.text_input("Semana", "")
curso = st.selectbox("Curso", ["1ro de Primaria", "2do de Primaria", "3ro de Primaria", 
                               "4to de Primaria", "5to de Primaria", "6to de Primaria"])
materia = st.text_input("Materia", "")
profesor = st.text_input("Profesor/a", "")

# Competencias
competencias_fundamentales = st.text_area("Competencias Fundamentales", "")

# Seleccionar fecha con un calendario
fecha = st.date_input("Fecha", date.today())
fecha_formateada = fecha.strftime("%A, %d de %B").capitalize()

# Datos para un d铆a
st.subheader(" Planificaci贸n Diaria")
competencias_especificas = st.text_area("Competencias Espec铆ficas", "")
contenido = st.text_area("Contenido", "")
indicadores_logros = st.text_area("Indicadores de Logros", "")
secuencia_actividades = st.text_area("Secuencia de Actividades", "")
ejes_tematicos = st.text_input("Ejes Tem谩ticos Transversales", "")
tipo_evaluacion = st.text_input("Tipo de Evaluaci贸n", "")
asignaciones = st.text_area("Asignaciones", "")

# Campo para subir im谩genes
st.subheader(" Recursos")
imagenes = st.file_uploader("Subir im谩genes (opcional)", type=["png", "jpg", "jpeg"], accept_multiple_files=True)

# Campo para ingresar links externos
links = st.text_area("Enlaces externos (opcional, uno por l铆nea)")

# Bot贸n para agregar la planificaci贸n del d铆a a la lista
if st.button("Agregar Planificaci贸n del D铆a"):
    planificacion_dia = {
        "fecha": fecha_formateada,
        "competencias_especificas": competencias_especificas,
        "contenido": contenido,
        "indicadores_logros": indicadores_logros,
        "secuencia_actividades": secuencia_actividades,
        "ejes_tematicos": ejes_tematicos,
        "tipo_evaluacion": tipo_evaluacion,
        "asignaciones": asignaciones,
        "imagenes": imagenes,
        "links": links.splitlines()
    }
    st.session_state.planificaciones.append(planificacion_dia)
    st.success(f"隆Planificaci贸n agregada para el d铆a {fecha_formateada}!")

# Mostrar planificaciones registradas
st.subheader(" Planificaciones Registradas")
if st.session_state.planificaciones:
    for idx, plan in enumerate(st.session_state.planificaciones, 1):
        st.write(f"**D铆a {idx}: {plan['fecha']}**")
        st.write(f"- **Competencias Espec铆ficas:** {plan['competencias_especificas']}")
        st.write(f"- **Contenido:** {plan['contenido']}")
        st.write(f"- **Indicadores de Logros:** {plan['indicadores_logros']}")
        st.write(f"- **Secuencia de Actividades:** {plan['secuencia_actividades']}")
        st.write(f"- **Ejes Tem谩ticos:** {plan['ejes_tematicos']}")
        st.write(f"- **Tipo de Evaluaci贸n:** {plan['tipo_evaluacion']}")
        st.write(f"- **Asignaciones:** {plan['asignaciones']}")
        st.write("---")
else:
    st.info("No hay planificaciones registradas a煤n.")

# Bot贸n para generar el documento consolidado
if st.button("Generar Documento Consolidado"):
    if st.session_state.planificaciones:
        doc = Document()
        doc.add_heading('Planificaci贸n Semanal Consolidada', level=1)
        doc.add_paragraph(f"Semana: {semana}        Curso: {curso}        Materia: {materia}        Profesor/a: {profesor}")
        
        doc.add_paragraph("Competencias Fundamentales:").bold = True
        doc.add_paragraph(competencias_fundamentales)
        
        # Cambiar la orientaci贸n a horizontal
        seccion = doc.sections[0]
        seccion.orientation = WD_ORIENTATION.LANDSCAPE
        seccion.page_width, seccion.page_height = seccion.page_height, seccion.page_width
        
        # Crear tabla
        tabla = doc.add_table(rows=1, cols=8)
        tabla.style = 'Table Grid'
        encabezados = ["Fecha", "Competencias Espec铆ficas", "Contenidos", "Indicadores de Logros", 
                       "Secuencia de Actividades", "Ejes Tem谩ticos Transversales", "Tipo de Evaluaci贸n", "Asignaciones"]
        
        hdr_cells = tabla.rows[0].cells
        for i, encabezado in enumerate(encabezados):
            hdr_cells[i].text = encabezado
            hdr_cells[i].paragraphs[0].runs[0].bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
        
        # Agregar filas con la informaci贸n de cada d铆a
        for plan in st.session_state.planificaciones:
            fila = tabla.add_row().cells
            fila[0].text = plan["fecha"]
            fila[1].text = plan["competencias_especificas"]
            fila[2].text = plan["contenido"]
            fila[3].text = plan["indicadores_logros"]
            fila[4].text = plan["secuencia_actividades"]
            fila[5].text = plan["ejes_tematicos"]
            fila[6].text = plan["tipo_evaluacion"]
            fila[7].text = plan["asignaciones"]
            
            # Agregar secci贸n de recursos con fecha
            if plan["imagenes"] or plan["links"]:
                doc.add_paragraph(f"Recursos para el d铆a {plan['fecha']}:", style='Heading 2')
                if plan["imagenes"]:
                    for img in plan["imagenes"]:
                        doc.add_paragraph(f"Imagen: {img.name}")
                        doc.add_picture(img, width=Pt(200))
                if plan["links"]:
                    for link in plan["links"]:
                        doc.add_paragraph(f"Enlace: {link}")
        
        doc.save('planificacion_consolidada.docx')
        st.success("隆Documento consolidado generado con 茅xito!")
        with open('planificacion_consolidada.docx', 'rb') as f:
            st.download_button("Descargar Planificaci贸n Consolidada", f, file_name="planificacion_consolidada.docx")
    else:
        st.warning("No hay planificaciones registradas para generar el documento.")
